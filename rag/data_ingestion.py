"""
Load the Word/PDF documents in rag_docs and prepare LangChain Documents.

The parent public-agent graph can use this module during indexing. It does not
define its own LangGraph because RAG is one node in the larger public-agent
workflow.
"""

from __future__ import annotations

from datetime import datetime
import logging
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover - reported clearly at runtime
    DocxDocument = None

try:
    from .chunking import DocumentChunker
except ImportError:  # Allows running as: python rag/data_ingestion.py
    from chunking import DocumentChunker

logger = logging.getLogger(__name__)

RAG_DIR = Path(__file__).resolve().parent
DEFAULT_DOCS_DIR = RAG_DIR / "rag_docs"

SUPPORTED_EXTENSIONS = {
    ".pdf": "PDF Document",
    ".docx": "Word Document",
    ".txt": "Text Document",
    ".md": "Markdown Document",
}


class DocumentLoader:
    """Load supported source files into LangChain Document objects."""

    def load_path(self, file_path: Path | str) -> List[Document]:
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")
        if not path.exists():
            raise FileNotFoundError(path)
        if path.stat().st_size == 0:
            raise ValueError(f"File is empty: {path}")

        if extension == ".pdf":
            documents = self._load_pdf(path)
        elif extension == ".docx":
            documents = self._load_docx(path)
        else:
            documents = self._load_text(path)

        for document in documents:
            document.metadata.update(self._base_metadata(path))

        return documents

    def _base_metadata(self, path: Path) -> Dict[str, Any]:
        return {
            "source": str(path),
            "filename": path.name,
            "extension": path.suffix.lower(),
            "file_type": SUPPORTED_EXTENSIONS[path.suffix.lower()],
            "file_size_bytes": path.stat().st_size,
            "loaded_at": datetime.now().isoformat(),
        }

    @staticmethod
    def _load_pdf(path: Path) -> List[Document]:
        loader = PyPDFLoader(str(path))
        documents = loader.load()
        if not documents:
            raise ValueError(f"No extractable text found in PDF: {path.name}")
        # logger.info("[INGESTION] Loaded %s PDF page document(s) from %s", len(documents), path.name)
        return documents

    @staticmethod
    def _load_docx(path: Path) -> List[Document]:
        if DocxDocument is None:
            raise RuntimeError("python-docx is required to load DOCX files.")

        docx = DocxDocument(str(path))
        parts: List[str] = []

        for paragraph in docx.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in docx.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        content = "\n\n".join(parts).strip()
        if not content:
            raise ValueError(f"No text found in DOCX: {path.name}")

        # logger.info("[INGESTION] Loaded DOCX document from %s", path.name)
        return [Document(page_content=content, metadata={})]

    @staticmethod
    def _load_text(path: Path) -> List[Document]:
        loader = TextLoader(str(path), encoding="utf-8")
        documents = loader.load()
        if not documents or not documents[0].page_content.strip():
            raise ValueError(f"No text found in file: {path.name}")
        return documents


class ColdMemoryIngestor:
    """Load documents from disk and chunk them for vector indexing."""

    def __init__(
        self,
        docs_dir: Path | str = DEFAULT_DOCS_DIR,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> None:
        self.docs_dir = Path(docs_dir)
        self.loader = DocumentLoader()
        self.chunker = DocumentChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    def ingest_file(self, file_path: Path | str) -> Tuple[List[Document], Dict[str, Any]]:
        path = Path(file_path)
        loaded_documents = self.loader.load_path(path)
        chunks = self.chunker.chunk_documents(loaded_documents)

        metadata = {
            "filename": path.name,
            "source": str(path),
            "document_count": len(loaded_documents),
            "chunk_count": len(chunks),
            "ingested_at": datetime.now().isoformat(),
            "memory_type": "cold",
        }
        return chunks, metadata

    def ingest_directory(self, docs_dir: Optional[Path | str] = None) -> Tuple[List[Document], List[Dict[str, Any]]]:
        root = Path(docs_dir) if docs_dir else self.docs_dir
        if not root.exists():
            raise FileNotFoundError(f"RAG docs folder does not exist: {root}")

        files = list(self.iter_supported_files(root))
        if not files:
            raise ValueError(f"No supported documents found in: {root}")

        all_chunks: List[Document] = []
        ingestion_report: List[Dict[str, Any]] = []

        for file_path in files:
            try:
                chunks, metadata = self.ingest_file(file_path)
                all_chunks.extend(chunks)
                ingestion_report.append(metadata)
            except Exception as exc:
                logger.exception("[INGESTION] Failed to ingest %s", file_path)
                ingestion_report.append(
                    {
                        "filename": file_path.name,
                        "source": str(file_path),
                        "chunk_count": 0,
                        "error": str(exc),
                    }
                )

        return all_chunks, ingestion_report

    @staticmethod
    def iter_supported_files(root: Path) -> Iterable[Path]:
        for path in sorted(root.rglob("*")):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                yield path


def load_rag_docs(
    docs_dir: Path | str = DEFAULT_DOCS_DIR,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
) -> Tuple[List[Document], List[Dict[str, Any]]]:
    """Load and chunk every supported file under rag/rag_docs."""
    ingestor = ColdMemoryIngestor(
        docs_dir=docs_dir,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    return ingestor.ingest_directory()


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(levelname)s:%(name)s:%(message)s")
    chunks, report = load_rag_docs()
    print(f"Loaded {len(chunks)} chunk(s) from {len(report)} file(s).")
    for item in report:
        print(item)
